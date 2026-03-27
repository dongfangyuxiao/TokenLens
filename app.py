import os, json, threading, secrets, re, zipfile, io, html as _html, subprocess, tempfile, shutil
from datetime import datetime, timedelta
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Form
from fastapi.responses import HTMLResponse, FileResponse, Response, PlainTextResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from typing import List, Optional
from apscheduler.schedulers.background import BackgroundScheduler
from urllib.parse import urlparse, urlunparse, quote
import database as db
import license_manager
from scanner import run_scan
from reporter import build_markdown_report, build_json_report
from notifier import test_channel as notifier_test_channel
import syslog_sender as syslog
from repo_sync import search_synced_code, normalize_repo_name
from clients.github    import test_connection as github_test
from clients.gitlab    import test_connection as gitlab_test
from clients.bitbucket import test_connection as bitbucket_test
from clients.gitee     import test_connection as gitee_test

# ── 登录鉴权 ──────────────────────────────────────────────────────
_sessions: dict = {}          # token -> {'username': str, 'expires_at': datetime, 'last_seen': datetime}
_login_attempts: dict = {}    # username -> {'count': int, 'lock_until': datetime|None}
_MAX_ATTEMPTS = 5
_LOCK_MINUTES = 10
_SESSION_TTL_MINUTES = max(10, int(os.getenv('SESSION_TTL_MINUTES', '720')))   # 默认 12 小时
_SESSION_IDLE_MINUTES = max(5, int(os.getenv('SESSION_IDLE_MINUTES', '120')))  # 默认 2 小时
_REPORTS_REQUIRE_AUTH = os.getenv('REPORTS_REQUIRE_AUTH', '1') == '1'

_http_bearer = HTTPBearer(auto_error=False)

def _check_password_strength(password: str):
    """返回错误信息，None 表示通过"""
    if len(password) < 8:
        return '密码长度不能少于 8 位'
    if not re.search(r'[a-z]', password):
        return '密码必须包含小写字母'
    if not re.search(r'[A-Z]', password):
        return '密码必须包含大写字母'
    if not re.search(r'\d', password):
        return '密码必须包含数字'
    return None

def _safe_http_url(url: str) -> str:
    u = (url or '').strip()
    if not u:
        return ''
    try:
        p = urlparse(u)
        if p.scheme in ('http', 'https'):
            return u
    except Exception:
        pass
    return ''

def _mask_secret(text: str, secret: str) -> str:
    if not text:
        return ''
    if not secret:
        return text
    try:
        masked = text.replace(secret, '***')
        masked = masked.replace(quote(secret, safe=''), '***')
        return masked
    except Exception:
        return text

def _purge_expired_sessions():
    now = datetime.now()
    dead = []
    for tok, sess in list(_sessions.items()):
        expires_at = sess.get('expires_at') if isinstance(sess, dict) else None
        last_seen = sess.get('last_seen') if isinstance(sess, dict) else None
        if not isinstance(expires_at, datetime) or not isinstance(last_seen, datetime):
            dead.append(tok)
            continue
        if now > expires_at or now - last_seen > timedelta(minutes=_SESSION_IDLE_MINUTES):
            dead.append(tok)
    for tok in dead:
        _sessions.pop(tok, None)

def _issue_session(username: str) -> str:
    now = datetime.now()
    token = secrets.token_hex(32)
    _sessions[token] = {
        'username': username,
        'expires_at': now + timedelta(minutes=_SESSION_TTL_MINUTES),
        'last_seen': now,
    }
    return token

def _auth_token_to_user(token: str, touch: bool = True) -> str:
    _purge_expired_sessions()
    if not token:
        return ''
    sess = _sessions.get(token)
    if not isinstance(sess, dict):
        return ''
    if touch:
        sess['last_seen'] = datetime.now()
    return sess.get('username', '')

def require_auth(credentials: HTTPAuthorizationCredentials = Depends(_http_bearer)):
    token = credentials.credentials if credentials else ''
    username = _auth_token_to_user(token, touch=True)
    if not username:
        raise HTTPException(status_code=401, detail='未授权，请先登录')
    return username

# ── 初始化 ────────────────────────────────────────────────────────
db.init_db()
db.mark_interrupted_scans()
syslog.reload(db.get_syslog_config())
app = FastAPI(title='春静企业代码安全平台')
app.add_middleware(CORSMiddleware, allow_origins=['*'], allow_methods=['*'], allow_headers=['*'])


@app.middleware('http')
async def security_headers_middleware(request, call_next):
    path = request.url.path or ''
    if _REPORTS_REQUIRE_AUTH and path.startswith('/reports/'):
        token = ''
        authz = request.headers.get('authorization', '')
        if authz.lower().startswith('bearer '):
            token = authz[7:].strip()
        if not token:
            token = (request.query_params.get('token') or '').strip()
        if not _auth_token_to_user(token, touch=True):
            return PlainTextResponse('未授权，请先登录后查看报告', status_code=401)

    response = await call_next(request)
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['Referrer-Policy'] = 'no-referrer'
    response.headers['Permissions-Policy'] = (
        'camera=(), microphone=(), geolocation=(), payment=(), usb=()'
    )
    if path.startswith('/reports/') and path.endswith('.html'):
        # 报告页安全头：限制资源来源，降低 XSS 利用风险
        response.headers['Content-Security-Policy'] = (
            "default-src 'self'; "
            "base-uri 'none'; "
            "object-src 'none'; "
            "frame-ancestors 'none'; "
            "img-src 'self' data:; "
            "style-src 'self' 'unsafe-inline'; "
            "script-src 'self' 'unsafe-inline'"
        )
    return response

BASE_URL = os.getenv('BASE_URL', 'http://localhost:8000')
REPORTS_DIR = os.getenv('REPORTS_DIR', 'reports')

# ── 定时任务 ──────────────────────────────────────────────────────
scheduler    = BackgroundScheduler()
_scan_lock   = threading.Lock()
_stop_event  = threading.Event()   # set → 请求停止
_pause_event = threading.Event()   # set → 运行中；clear → 暂停中
_pause_event.set()
_scan_progress_lock = threading.Lock()
_scan_progress = {'percent': 0, 'phase': 'idle', 'message': '', 'updated_at': ''}

def _set_scan_progress(percent=None, phase=None, message=None):
    with _scan_progress_lock:
        if percent is not None:
            _scan_progress['percent'] = max(0, min(100, int(percent)))
        if phase is not None:
            _scan_progress['phase'] = phase
        if message is not None:
            _scan_progress['message'] = message
        _scan_progress['updated_at'] = datetime.now().isoformat()

def _get_scan_progress():
    with _scan_progress_lock:
        return dict(_scan_progress)

_WEEKDAY_CN = ['周一','周二','周三','周四','周五','周六','周日']
_LICENSE_WARN_DAYS = 30
_FEATURE_LABELS = {
    'poison_scan': '投毒检测',
    'incremental_audit': '增量代码审计',
    'full_audit': '全量代码审计',
    'instant_analysis': '即时分析',
}


def _get_license_status():
    cfg = db.get_app_config()
    license_key = (cfg.get('license_key') or '').strip()
    enforce_enabled = cfg.get('license_enforce_enabled', '0') == '1'
    instance_id = license_manager.get_instance_id()
    result = license_manager.verify_license(
        license_key,
        expected_product=license_manager.DEFAULT_PRODUCT,
        expected_machine_id=instance_id,
    )
    payload = result.get('payload') or {}
    expires_at = payload.get('expires_at', '')
    expires_in_days = None
    warning = ''
    if expires_at:
        try:
            expires_dt = license_manager._parse_iso8601(expires_at)
            if expires_dt:
                delta = expires_dt - datetime.now(expires_dt.tzinfo)
                expires_in_days = max(delta.days, 0)
                if result.get('valid') and expires_in_days <= _LICENSE_WARN_DAYS:
                    warning = f'授权将在 {expires_in_days} 天后到期'
        except Exception:
            expires_in_days = None
    return {
        'configured': bool(license_key),
        'enforce_enabled': enforce_enabled,
        'instance_id': instance_id,
        'valid': result.get('valid', False),
        'state': result.get('state', 'missing'),
        'message': result.get('message', '未配置授权码'),
        'product': payload.get('product', license_manager.DEFAULT_PRODUCT),
        'customer': payload.get('customer', ''),
        'issued_at': payload.get('issued_at', ''),
        'expires_at': expires_at,
        'expires_in_days': expires_in_days,
        'warning': warning,
        'features': payload.get('features', []),
        'machine_id': payload.get('machine_id', ''),
        'metadata': payload.get('metadata', {}),
    }


def _license_allows_feature(status: dict, feature: str) -> bool:
    features = status.get('features') or []
    if not features:
        return True
    return feature in features


def _require_license_if_enabled(feature: str = ''):
    status = _get_license_status()
    if status['enforce_enabled'] and not status['valid']:
        raise HTTPException(403, f'产品授权校验失败：{status["message"]}')
    if status['enforce_enabled'] and feature and not _license_allows_feature(status, feature):
        raise HTTPException(403, f'当前授权未开通：{_FEATURE_LABELS.get(feature, feature)}')
    return status

def _run_in_thread(scan_type: str, llm_profile_id=None):
    """在新线程中执行扫描，持有 _scan_lock。"""
    if not _scan_lock.acquire(blocking=False):
        print('[scheduler] 上次扫描未结束，跳过')
        return
    _stop_event.clear()
    _pause_event.set()
    _set_scan_progress(0, 'starting', f'{scan_type} starting')
    try:
        run_scan(BASE_URL, scan_type=scan_type,
                 stop_event=_stop_event, pause_event=_pause_event,
                 llm_profile_id=llm_profile_id,
                 progress_cb=_set_scan_progress)
    finally:
        _scan_lock.release()
        _set_scan_progress(0, 'idle', '')

def _reschedule():
    scheduler.remove_all_jobs()
    for s in db.get_scan_schedules():
        if not s['enabled']:
            continue
        if s['type'] not in {'poison', 'incremental_audit', 'full_audit'}:
            continue
        scan_type      = s['type']
        llm_profile_id = s.get('llm_profile_id')
        h  = s['hour']
        m  = s['minute']
        wd = s.get('weekday')
        fn = lambda st=scan_type, lpid=llm_profile_id: _run_in_thread(st, lpid)
        if h == -1:
            scheduler.add_job(fn, 'cron', hour='*', minute=m, id=f"sched_{s['id']}")
            print(f"[scheduler] {scan_type} #{s['id']}: 每小时第 {m:02d} 分")
        elif wd is not None:
            scheduler.add_job(fn, 'cron', day_of_week=wd, hour=h, minute=m,
                              id=f"sched_{s['id']}")
            print(f"[scheduler] {scan_type} #{s['id']}: 每{_WEEKDAY_CN[wd]} {h:02d}:{m:02d}")
        else:
            scheduler.add_job(fn, 'cron', hour=h, minute=m, id=f"sched_{s['id']}")
            print(f"[scheduler] {scan_type} #{s['id']}: 每天 {h:02d}:{m:02d}")

_reschedule()
scheduler.start()

# ── 静态文件 ──────────────────────────────────────────────────────
os.makedirs(REPORTS_DIR, exist_ok=True)
os.makedirs('static', exist_ok=True)
app.mount('/reports', StaticFiles(directory=REPORTS_DIR), name='reports')

@app.get('/', response_class=HTMLResponse)
def index():
    path = os.path.join(os.path.dirname(__file__), 'static', 'index.html')
    return open(path, encoding='utf-8').read()

# ── 登录 / 登出 ───────────────────────────────────────────────────
class LoginIn(BaseModel):
    username: str
    password: str

@app.post('/api/login')
def login(body: LoginIn):
    now = datetime.now()
    username = body.username
    attempt = _login_attempts.get(username, {'count': 0, 'lock_until': None})

    # 检查是否处于锁定期
    if attempt['lock_until'] and now < attempt['lock_until']:
        remaining = int((attempt['lock_until'] - now).total_seconds() / 60) + 1
        syslog.send('warning', 'AUTH', f'登录被拒 (账户锁定): {username}')
        raise HTTPException(403, f'账户已锁定，请 {remaining} 分钟后重试')

    # 锁定已到期则重置
    if attempt['lock_until'] and now >= attempt['lock_until']:
        attempt = {'count': 0, 'lock_until': None}

    if not db.verify_admin(username, body.password):
        attempt['count'] += 1
        if attempt['count'] >= _MAX_ATTEMPTS:
            attempt['lock_until'] = now + timedelta(minutes=_LOCK_MINUTES)
            attempt['count'] = 0
            _login_attempts[username] = attempt
            syslog.send('warning', 'AUTH', f'账户锁定 (连续失败 {_MAX_ATTEMPTS} 次): {username}')
            raise HTTPException(403, f'密码错误次数过多，账户已锁定 {_LOCK_MINUTES} 分钟')
        _login_attempts[username] = attempt
        left = _MAX_ATTEMPTS - attempt['count']
        syslog.send('warning', 'AUTH', f'登录失败 (剩余 {left} 次): {username}')
        raise HTTPException(401, f'用户名或密码错误，还有 {left} 次机会')

    # 登录成功
    _login_attempts.pop(username, None)
    token = _issue_session(username)
    syslog.send('info', 'AUTH', f'登录成功: {username}')
    return {'token': token, 'username': username}

@app.post('/api/logout')
def logout(credentials: HTTPAuthorizationCredentials = Depends(_http_bearer)):
    if credentials and credentials.credentials in _sessions:
        sess = _sessions.pop(credentials.credentials)
        user = sess.get('username') if isinstance(sess, dict) else str(sess)
        syslog.send('info', 'AUTH', f'登出: {user}')
    return {'ok': True}

@app.get('/api/me')
def me(username: str = Depends(require_auth)):
    return {'username': username}

# ── 用户管理 API ──────────────────────────────────────────────────
class UserIn(BaseModel):
    username: str
    password: str

class PasswordIn(BaseModel):
    password: str

@app.get('/api/admin-users')
def list_admin_users(_: str = Depends(require_auth)):
    return db.get_admin_users()

@app.post('/api/admin-users')
def create_admin_user(body: UserIn, _: str = Depends(require_auth)):
    if not body.username or len(body.username.strip()) < 2:
        raise HTTPException(400, '用户名至少 2 个字符')
    err = _check_password_strength(body.password)
    if err:
        raise HTTPException(400, err)
    if not db.add_admin_user(body.username.strip(), body.password):
        raise HTTPException(409, '用户名已存在')
    syslog.send('info', 'ADMIN', f'添加用户: {body.username.strip()} by {_}')
    return {'ok': True}

@app.delete('/api/admin-users/{username}')
def delete_admin_user(username: str, current: str = Depends(require_auth)):
    if username == current:
        raise HTTPException(400, '不能删除当前登录的账户')
    if not db.delete_admin_user(username):
        if not any(u['username'] == username for u in db.get_admin_users()):
            raise HTTPException(404, '用户不存在')
        raise HTTPException(400, '至少保留一个管理员账户')
    syslog.send('info', 'ADMIN', f'删除用户: {username} by {current}')
    return {'ok': True}

@app.patch('/api/admin-users/{username}/password')
def change_password(username: str, body: PasswordIn, current: str = Depends(require_auth)):
    err = _check_password_strength(body.password)
    if err:
        raise HTTPException(400, err)
    if not db.update_admin_password(username, body.password):
        raise HTTPException(404, '用户不存在')
    syslog.send('info', 'ADMIN', f'修改密码: {username} by {current}')
    return {'ok': True}

# ── 账户 API ──────────────────────────────────────────────────────
class AccountIn(BaseModel):
    platform: str
    name    : str
    token   : str
    url     : str = ''
    owner   : str = ''

@app.get('/api/accounts')
def list_accounts(_: str = Depends(require_auth)):
    rows = db.get_accounts()
    # 隐藏 token
    for r in rows:
        t = r.get('token', '')
        r['token_masked'] = t[:6] + '****' + t[-4:] if len(t) > 10 else '****'
    return rows

@app.post('/api/accounts')
def create_account(body: AccountIn, _: str = Depends(require_auth)):
    db.add_account(body.platform, body.name, body.token, body.url, body.owner)
    return {'ok': True}

@app.delete('/api/accounts/{aid}')
def remove_account(aid: int, _: str = Depends(require_auth)):
    db.delete_account(aid)
    return {'ok': True}

class AccountTestIn(BaseModel):
    platform: str
    token   : str
    url     : str = ''

@app.post('/api/accounts/test')
def test_account(body: AccountTestIn, _: str = Depends(require_auth)):
    platform = body.platform
    token    = body.token
    raw_url  = body.url.strip()
    if raw_url and not _safe_http_url(raw_url):
        raise HTTPException(400, '服务器地址格式不正确，仅支持 http/https')
    url      = raw_url or 'https://gitlab.com'
    if platform == 'github':
        err, user = github_test(token)
    elif platform == 'bitbucket':
        err, user = bitbucket_test(token)
    elif platform == 'gitee':
        err, user = gitee_test(token)
    elif platform in ('tgit', 'codeup'):
        _default_urls = {'tgit': 'https://git.code.tencent.com', 'codeup': 'https://codeup.aliyun.com'}
        base = body.url.strip() or _default_urls[platform]
        err, user = gitlab_test(token, base)
    else:  # gitlab
        err, user = gitlab_test(token, url)
    if err:
        raise HTTPException(400, f'连接失败: {_mask_secret(str(err), token)}')
    return {'ok': True, 'msg': f'连接成功，当前用户: {user}'}

# ── 通知渠道 API ──────────────────────────────────────────────────
class ChannelIn(BaseModel):
    name  : str
    type  : str
    config: dict

@app.get('/api/channels')
def list_channels(_: str = Depends(require_auth)):
    return db.get_channels()

@app.post('/api/channels')
def create_channel(body: ChannelIn, _: str = Depends(require_auth)):
    db.add_channel(body.name, body.type, body.config)
    return {'ok': True}

@app.patch('/api/channels/{cid}')
def toggle_channel(cid: int, body: dict, _: str = Depends(require_auth)):
    db.toggle_channel(cid, body.get('enabled', True))
    return {'ok': True}

class ChannelUpdateIn(BaseModel):
    name  : str
    config: dict

@app.put('/api/channels/{cid}')
def update_channel(cid: int, body: ChannelUpdateIn, _: str = Depends(require_auth)):
    if not body.name.strip():
        raise HTTPException(400, '渠道名称不能为空')
    db.update_channel(cid, body.name.strip(), body.config)
    return {'ok': True}

@app.delete('/api/channels/{cid}')
def remove_channel(cid: int, _: str = Depends(require_auth)):
    db.delete_channel(cid)
    return {'ok': True}

@app.post('/api/channels/{cid}/test')
def test_channel_api(cid: int, _: str = Depends(require_auth)):
    channels = db.get_channels()
    ch = next((c for c in channels if c['id'] == cid), None)
    if not ch:
        raise HTTPException(404, '渠道不存在')
    cfg = json.loads(ch.get('config', '{}'))
    err = notifier_test_channel(ch['type'], cfg)
    if err:
        raise HTTPException(400, f'发送失败: {err}')
    return {'ok': True, 'msg': f'测试消息已发送至 {ch["name"]}'}

# ── LLM 配置 API ─────────────────────────────────────────────────
# ── LLM 配置列表 API ──────────────────────────────────────────────
class LLMProfileIn(BaseModel):
    name    : str
    provider: str
    model   : str = ''
    api_key : str = ''
    base_url: str = ''

@app.get('/api/llm-profiles')
def list_llm_profiles(_: str = Depends(require_auth)):
    return db.get_llm_profiles()

@app.post('/api/llm-profiles')
def add_llm_profile(body: LLMProfileIn, _: str = Depends(require_auth)):
    if not body.name.strip() or not body.provider.strip():
        raise HTTPException(400, '名称和提供商不能为空')
    pid = db.add_llm_profile(body.name.strip(), body.provider, body.model, body.api_key, body.base_url)
    return {'ok': True, 'id': pid}

@app.put('/api/llm-profiles/{pid}')
def update_llm_profile(pid: int, body: LLMProfileIn, _: str = Depends(require_auth)):
    if not db.get_llm_profile(pid):
        raise HTTPException(404, '配置不存在')
    db.update_llm_profile(pid, body.name.strip(), body.provider, body.model, body.api_key, body.base_url)
    return {'ok': True}

@app.delete('/api/llm-profiles/{pid}')
def delete_llm_profile(pid: int, _: str = Depends(require_auth)):
    db.delete_llm_profile(pid)
    return {'ok': True}

@app.post('/api/llm-profiles/{pid}/test')
def test_llm_profile(pid: int, _: str = Depends(require_auth)):
    from analyzer import build_llm_caller
    profile = db.get_llm_profile(pid)
    if not profile:
        raise HTTPException(404, '配置不存在')
    llm_cfg = {'provider': profile['provider'], 'model': profile['model'],
               'api_key': profile['api_key'], 'base_url': profile['base_url']}
    call_fn = build_llm_caller(llm_cfg)
    if not call_fn:
        raise HTTPException(400, '未配置 API Key')
    try:
        reply = call_fn('Reply with exactly one word: OK')
        return {'ok': True, 'reply': reply.strip()}
    except Exception as e:
        raise HTTPException(502, str(e))

_SENSITIVE_LLM_KEYS = {'api_key'}

class LLMConfig(BaseModel):
    provider : str = ''
    model    : str = ''
    api_key  : str = ''
    base_url : str = ''

@app.get('/api/llm-config')
def get_llm(_: str = Depends(require_auth)):
    cfg = db.get_llm_config()
    # 仅返回新版字段，隐藏旧版兼容字段
    cfg = {k: cfg.get(k, '') for k in ('provider', 'model', 'api_key', 'base_url')}
    def mask(k, v):
        if k in _SENSITIVE_LLM_KEYS:
            return v[:8] + '****' if len(v) > 8 else ('****' if v else '')
        return v
    return {k: mask(k, v) for k, v in cfg.items()}

@app.post('/api/llm-config')
def save_llm(body: LLMConfig, _: str = Depends(require_auth)):
    if body.provider is not None:
        db.set_llm_config('provider', body.provider)
    if body.model is not None:
        db.set_llm_config('model', body.model)
    if body.base_url is not None:
        db.set_llm_config('base_url', body.base_url)
    if body.api_key and '****' not in body.api_key:
        db.set_llm_config('api_key', body.api_key)
    return {'ok': True}

@app.post('/api/llm-config/test')
def test_llm(_: str = Depends(require_auth)):
    from analyzer import build_llm_caller
    llm_cfg = db.get_llm_config()
    call_fn = build_llm_caller(llm_cfg)
    if not call_fn:
        raise HTTPException(400, '未配置 LLM，请先填写提供商和 API Key')
    try:
        reply = call_fn('Reply with exactly one word: OK')
        return {'ok': True, 'reply': reply.strip()}
    except Exception as e:
        raise HTTPException(502, str(e))

# ── 应用设置 API ──────────────────────────────────────────────────
@app.get('/api/settings')
def get_settings(_: str = Depends(require_auth)):
    return db.get_app_config()

@app.post('/api/settings')
def save_settings(body: dict, _: str = Depends(require_auth)):
    for k, v in body.items():
        db.set_app_config(k, v)
    _reschedule()
    return {'ok': True}


@app.get('/api/license-status')
def get_license_status(_: str = Depends(require_auth)):
    return _get_license_status()


class LicenseConfigIn(BaseModel):
    license_key: str = ''
    replace_license_key: bool = False
    enforce_enabled: bool = False

class LicenseGenerateIn(BaseModel):
    customer: str
    expires_at: str
    machine_id: str = ''
    features: List[str] = []
    product: str = license_manager.DEFAULT_PRODUCT
    metadata: dict = {}


@app.post('/api/license-config')
def save_license_config(body: LicenseConfigIn, _: str = Depends(require_auth)):
    if body.replace_license_key:
        db.set_app_config('license_key', body.license_key.strip())
    db.set_app_config('license_enforce_enabled', '1' if body.enforce_enabled else '0')
    return {'ok': True, 'status': _get_license_status()}

@app.get('/api/license/machine-file')
def download_machine_file(_: str = Depends(require_auth)):
    payload = {
        'product': license_manager.DEFAULT_PRODUCT,
        'instance_id': license_manager.get_instance_id(),
        'generated_at': datetime.now().isoformat(),
    }
    content = json.dumps(payload, ensure_ascii=False, indent=2)
    fname = f'machine_id_{payload["instance_id"]}.json'
    return Response(
        content=content.encode('utf-8'),
        media_type='application/json; charset=utf-8',
        headers={'Content-Disposition': f'attachment; filename="{fname}"'}
    )

@app.post('/api/license/upload-file')
async def upload_license_file(
    file: UploadFile = File(...),
    enforce_enabled: Optional[bool] = Form(None),
    _: str = Depends(require_auth),
):
    raw = await file.read()
    if not raw:
        raise HTTPException(400, '授权文件为空')
    if len(raw) > 200_000:
        raise HTTPException(400, '授权文件过大')

    text = raw.decode('utf-8', errors='replace').strip()
    if not text:
        raise HTTPException(400, '授权文件内容为空')

    license_key = ''
    if text.startswith('{'):
        try:
            data = json.loads(text)
            license_key = (data.get('license_key') or data.get('token') or '').strip()
        except Exception:
            raise HTTPException(400, '授权文件不是有效 JSON')
    else:
        license_key = text
    if not license_key:
        raise HTTPException(400, '授权文件中未找到 license_key')

    db.set_app_config('license_key', license_key)
    if enforce_enabled is not None:
        db.set_app_config('license_enforce_enabled', '1' if enforce_enabled else '0')
    return {'ok': True, 'status': _get_license_status()}

@app.post('/api/license/generate-file')
def generate_license_file(body: LicenseGenerateIn, _: str = Depends(require_auth)):
    customer = (body.customer or '').strip()
    if not customer:
        raise HTTPException(400, '客户名称不能为空')
    expires_at = (body.expires_at or '').strip()
    if not license_manager._parse_iso8601(expires_at):
        raise HTTPException(400, '过期时间格式不正确，请使用 ISO8601 格式')

    machine_id = (body.machine_id or '').strip() or license_manager.get_instance_id()
    features = [str(x).strip() for x in (body.features or []) if str(x).strip()]
    payload = license_manager.build_payload(
        customer=customer,
        expires_at=expires_at,
        product=(body.product or license_manager.DEFAULT_PRODUCT).strip() or license_manager.DEFAULT_PRODUCT,
        features=features,
        machine_id=machine_id,
        metadata=body.metadata or {},
    )
    license_key = license_manager.generate_license(payload)
    out = {**payload, 'license_key': license_key}
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_customer = re.sub(r'[^0-9A-Za-z._-]+', '_', customer)[:32] or 'customer'
    fname = f'license_{safe_customer}_{machine_id[:8]}_{ts}.json'
    return Response(
        content=json.dumps(out, ensure_ascii=False, indent=2).encode('utf-8'),
        media_type='application/json; charset=utf-8',
        headers={'Content-Disposition': f'attachment; filename="{fname}"'}
    )

# ── 扫描时间表 API ────────────────────────────────────────────────
_VALID_TYPES = {'poison', 'incremental_audit', 'full_audit'}

class ScanScheduleIn(BaseModel):
    type          : str
    hour          : int          # -1 = 每小时
    minute        : int = 0
    weekday       : int = None   # None=每天, 0-6=指定星期
    label         : str = ''
    llm_profile_id: int = None

@app.get('/api/scan-schedules')
def list_scan_schedules(_: str = Depends(require_auth)):
    return db.get_scan_schedules()

@app.post('/api/scan-schedules')
def add_scan_schedule(body: ScanScheduleIn, _: str = Depends(require_auth)):
    if body.type not in _VALID_TYPES:
        raise HTTPException(400, f'无效类型，可选: {", ".join(_VALID_TYPES)}')
    if body.hour != -1 and not (0 <= body.hour <= 23):
        raise HTTPException(400, '小时必须在 0-23 之间，或 -1 表示每小时')
    if not (0 <= body.minute <= 59):
        raise HTTPException(400, '分钟必须在 0-59 之间')
    if body.weekday is not None and not (0 <= body.weekday <= 6):
        raise HTTPException(400, '星期必须在 0-6 之间')
    sid = db.add_scan_schedule(body.type, body.hour, body.minute,
                               body.weekday, body.label.strip(), body.llm_profile_id)
    _reschedule()
    return {'ok': True, 'id': sid}

@app.patch('/api/scan-schedules/{sid}')
def toggle_scan_schedule(sid: int, body: dict, _: str = Depends(require_auth)):
    db.toggle_scan_schedule(sid, body.get('enabled', True))
    _reschedule()
    return {'ok': True}

@app.delete('/api/scan-schedules/{sid}')
def delete_scan_schedule(sid: int, _: str = Depends(require_auth)):
    db.delete_scan_schedule(sid)
    _reschedule()
    return {'ok': True}

# ── 扫描 API ──────────────────────────────────────────────────────
_SCAN_TYPE_LABELS = {
    'poison'           : '投毒检测',
    'incremental_audit': '增量代码审计',
    'full_audit'       : '全量代码审计',
}

class ScanTrigger(BaseModel):
    scan_type     : str  = 'incremental_audit'
    llm_profile_id: int  = None
    selected_sources: List[str] = []
    # 旧版兼容
    full_scan     : bool = False

@app.post('/api/scan/trigger')
def trigger_scan(body: ScanTrigger = ScanTrigger(), operator: str = Depends(require_auth)):
    scan_type = body.scan_type
    if body.full_scan and scan_type == 'incremental_audit':
        scan_type = 'full_audit'
    if scan_type not in _SCAN_TYPE_LABELS:
        raise HTTPException(400, f'无效类型，可选: {", ".join(_SCAN_TYPE_LABELS)}')
    feature_key = 'poison_scan' if scan_type == 'poison' else scan_type
    _require_license_if_enabled(feature_key)
    if not _scan_lock.acquire(blocking=False):
        raise HTTPException(400, '扫描正在进行中')
    _stop_event.clear()
    _pause_event.set()
    _set_scan_progress(0, 'starting', f'{scan_type} starting')
    label = _SCAN_TYPE_LABELS[scan_type]
    syslog.send('info', 'SCAN', f'{label}由 {operator} 触发')
    llm_profile_id = body.llm_profile_id
    selected_sources = body.selected_sources or []
    def _run():
        try:
            run_scan(BASE_URL, scan_type=scan_type,
                     stop_event=_stop_event, pause_event=_pause_event,
                     manual=True, llm_profile_id=llm_profile_id,
                     selected_sources=selected_sources,
                     progress_cb=_set_scan_progress)
        finally:
            _scan_lock.release()
            _set_scan_progress(0, 'idle', '')
    threading.Thread(target=_run, daemon=True).start()
    return {'ok': True, 'msg': f'{label}已启动'}

@app.post('/api/scan/stop')
def stop_scan(operator: str = Depends(require_auth)):
    if _scan_lock.acquire(blocking=False):
        _scan_lock.release()
        raise HTTPException(400, '当前没有正在进行的扫描')
    _stop_event.set()
    _pause_event.set()   # 解除暂停，让线程能检测到 stop
    syslog.send('info', 'SCAN', f'扫描被 {operator} 手动停止')
    return {'ok': True, 'msg': '已发送停止信号'}

@app.post('/api/scan/pause')
def pause_scan(operator: str = Depends(require_auth)):
    if _scan_lock.acquire(blocking=False):
        _scan_lock.release()
        raise HTTPException(400, '当前没有正在进行的扫描')
    _pause_event.clear()
    # 更新 DB 中正在运行的扫描状态为 paused
    running = [s for s in db.get_scans(limit=5) if s['status'] == 'running']
    if running:
        db.update_scan_status(running[0]['id'], 'paused')
    syslog.send('info', 'SCAN', f'扫描被 {operator} 暂停')
    return {'ok': True, 'msg': '扫描已暂停'}

@app.post('/api/scan/resume')
def resume_scan(operator: str = Depends(require_auth)):
    if _scan_lock.acquire(blocking=False):
        _scan_lock.release()
        raise HTTPException(400, '当前没有正在进行的扫描')
    _pause_event.set()
    # 更新 DB 中暂停的扫描状态为 running
    paused = [s for s in db.get_scans(limit=5) if s['status'] == 'paused']
    if paused:
        db.update_scan_status(paused[0]['id'], 'running')
    syslog.send('info', 'SCAN', f'扫描被 {operator} 恢复')
    return {'ok': True, 'msg': '扫描已恢复'}

@app.get('/api/scans')
def list_scans(_: str = Depends(require_auth)):
    return db.get_scans()

@app.delete('/api/scans/{scan_id}')
def delete_scan(scan_id: int, _: str = Depends(require_auth)):
    # 禁止删除正在运行的扫描
    row = db.get_scans(limit=200)
    target = next((s for s in row if s['id'] == scan_id), None)
    if not target:
        raise HTTPException(404, '扫描记录不存在')
    if target['status'] in ('running', 'paused'):
        raise HTTPException(400, '无法删除正在运行或暂停中的扫描，请先停止')
    db.delete_scan(scan_id)
    return {'ok': True}

@app.post('/api/scans/{scan_id}/rerun')
def rerun_scan(scan_id: int, operator: str = Depends(require_auth)):
    if not _scan_lock.acquire(blocking=False):
        raise HTTPException(400, '扫描正在进行中，请等待结束后重新扫描')
    row = db.get_scans(limit=200)
    target = next((s for s in row if s['id'] == scan_id), None)
    if not target:
        _scan_lock.release()
        raise HTTPException(404, '扫描记录不存在')
    scan_type = target.get('scan_type') or 'incremental_audit'
    # 兼容旧记录
    if scan_type == 'full':
        scan_type = 'full_audit'
    elif scan_type == 'incremental':
        scan_type = 'incremental_audit'
    feature_key = 'poison_scan' if scan_type == 'poison' else scan_type
    _require_license_if_enabled(feature_key)
    _stop_event.clear()
    _pause_event.set()
    _set_scan_progress(0, 'starting', f'{scan_type} rerun starting')
    label          = _SCAN_TYPE_LABELS.get(scan_type, scan_type)
    llm_profile_id = target.get('llm_profile_id')
    syslog.send('info', 'SCAN', f'{label}（重新运行 #{scan_id}）由 {operator} 触发')
    def _run():
        try:
            run_scan(BASE_URL, scan_type=scan_type,
                     stop_event=_stop_event, pause_event=_pause_event,
                     manual=True, llm_profile_id=llm_profile_id,
                     progress_cb=_set_scan_progress)
        finally:
            _scan_lock.release()
            _set_scan_progress(0, 'idle', '')
    threading.Thread(target=_run, daemon=True).start()
    return {'ok': True, 'msg': f'{label}已重新启动'}

@app.get('/api/scans/{scan_id}/findings')
def get_findings(scan_id: int, _: str = Depends(require_auth)):
    rows = db.get_scan_findings(scan_id)
    for r in rows:
        r['commit_url'] = _safe_http_url(r.get('commit_url', ''))
    return rows

@app.get('/api/scans/{scan_id}/logs')
def get_scan_logs(scan_id: int, _: str = Depends(require_auth)):
    return db.get_scan_logs(scan_id)

def _build_docx_report(scan_results):
    """生成 Word 格式审计报告，返回 bytes。"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    SEV_ZH = {'critical': '严重', 'high': '高危', 'medium': '中危', 'low': '低危', 'info': '信息'}
    SEV_COLOR = {'critical': 'FF4D4F', 'high': 'FA8C16', 'medium': 'D4B106', 'low': '52C41A', 'info': '1677FF'}
    doc = Document()
    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
    doc.add_heading('代码安全审计报告', 0)
    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    total = sum(len(e.get('findings', [])) for e in scan_results)
    doc.add_paragraph(f'共发现漏洞：{total} 个，涉及仓库：{len(scan_results)} 个')
    doc.add_paragraph()
    for entry in scan_results:
        doc.add_heading(f'仓库：{entry["repo"]}', level=1)
        meta = doc.add_paragraph()
        meta.add_run('提交 SHA：').bold = True
        meta.add_run(entry.get('commit_sha', '-')[:12])
        meta.add_run('　作者：').bold = True
        meta.add_run(entry.get('author', '-'))
        if entry.get('committed_at'):
            meta.add_run(f'　时间：{entry["committed_at"]}')
        findings = entry.get('findings', [])
        if not findings:
            doc.add_paragraph('（本提交无发现漏洞）')
            continue
        for finding in findings:
            sev = finding.get('severity', 'info')
            sev_zh = SEV_ZH.get(sev, sev)
            h = doc.add_heading(f'[{sev_zh}] {finding.get("title", "未知漏洞")}', level=2)
            run = h.runs[0] if h.runs else None
            if run:
                color = SEV_COLOR.get(sev, '1677FF')
                run.font.color.rgb = RGBColor(int(color[:2],16), int(color[2:4],16), int(color[4:],16))
            if finding.get('filename'):
                p = doc.add_paragraph()
                p.add_run('文件：').bold = True
                p.add_run(finding['filename'])
            if finding.get('description'):
                doc.add_paragraph(finding['description'])
            if finding.get('recommendation'):
                p2 = doc.add_paragraph()
                p2.add_run('修复建议：').bold = True
                p2.add_run(finding['recommendation'])
            doc.add_paragraph()
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

@app.get('/api/scans/{scan_id}/export')
def export_scan(scan_id: int, format: str = 'json', _: str = Depends(require_auth)):
    """导出扫描结果为 json / markdown / docx 格式。"""
    if format not in ('json', 'markdown', 'docx'):
        raise HTTPException(400, '无效格式，可选：json / markdown / docx')
    scan = db.get_scan(scan_id)
    if not scan:
        raise HTTPException(404, '扫描记录不存在')
    findings = db.get_scan_findings(scan_id)
    # 重组为 build_*_report 所需的 scan_results 格式
    commit_map: dict = {}
    for f in findings:
        key = (f['commit_sha'], f['repo'])
        if key not in commit_map:
            commit_map[key] = {
                'repo'        : f['repo'],
                'source'      : '',
                'commit_sha'  : f['commit_sha'],
                'commit_url'  : f['commit_url'],
                'author'      : f['author'],
                'message'     : '',
                'committed_at': f['committed_at'],
                'files'       : [],
                'findings'    : [],
            }
        commit_map[key]['findings'].append(f)
    scan_results = list(commit_map.values())
    if format == 'json':
        data = build_json_report(scan_results)
        return Response(
            content=json.dumps(data, ensure_ascii=False, indent=2),
            media_type='application/json',
            headers={'Content-Disposition': f'attachment; filename="scan_{scan_id}.json"'}
        )
    elif format == 'docx':
        content = _build_docx_report(scan_results)
        return Response(
            content=content,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="scan_{scan_id}.docx"'}
        )
    else:
        md = build_markdown_report(scan_results)
        return Response(
            content=md,
            media_type='text/markdown; charset=utf-8',
            headers={'Content-Disposition': f'attachment; filename="scan_{scan_id}.md"'}
        )

@app.get('/api/findings/{finding_id}')
def get_finding(finding_id: int, _: str = Depends(require_auth)):
    """获取单条漏洞详情，含仓库负责人信息。"""
    f = db.get_finding_detail(finding_id)
    if not f:
        raise HTTPException(404, '漏洞不存在')
    f['commit_url'] = _safe_http_url(f.get('commit_url', ''))
    return f

class FindingStatusIn(BaseModel):
    status: str

@app.patch('/api/findings/{finding_id}/status')
def update_finding_status(finding_id: int, body: FindingStatusIn, _: str = Depends(require_auth)):
    if not db.update_finding_status(finding_id, body.status):
        raise HTTPException(400, '无效的状态值或漏洞不存在')
    return {'ok': True}

@app.get('/api/stats')
def get_stats_api(period: str = 'day', repo: str = 'all', _: str = Depends(require_auth)):
    if period not in ('day', 'week', 'year'):
        raise HTTPException(400, '无效的统计周期，可选：day / week / year')
    return db.get_stats(period, repo)

@app.get('/api/repos')
def list_repos(_: str = Depends(require_auth)):
    return db.get_repos()

@app.get('/api/repo-sync-status')
def repo_sync_status(_: str = Depends(require_auth)):
    return db.get_repo_sync_status()

@app.get('/api/components')
def list_components(keyword: str = '', repo: str = '', source: str = '',
                    ecosystem: str = '', limit: int = 500,
                    _: str = Depends(require_auth)):
    return db.query_components(keyword=keyword, repo=repo, source=source,
                               ecosystem=ecosystem, limit=limit)

@app.get('/api/components/summary')
def component_summary(_: str = Depends(require_auth)):
    return db.get_component_summary()

class EmergencyDependencyIn(BaseModel):
    keyword: str
    exact: bool = False
    source: str = ''
    repo: str = ''
    ecosystem: str = ''
    limit: int = 500

@app.post('/api/emergency/dependency-check')
def emergency_dependency_check(body: EmergencyDependencyIn, _: str = Depends(require_auth)):
    keyword = body.keyword.strip()
    if not keyword:
        raise HTTPException(400, '依赖关键字不能为空')
    rows = db.query_components(
        keyword='' if body.exact else keyword,
        repo=normalize_repo_name(body.repo),
        source=body.source.strip(),
        ecosystem=body.ecosystem.strip(),
        limit=body.limit,
    )
    if body.exact:
        rows = [r for r in rows if (r.get('component', '').lower() == keyword.lower())]
    return {
        'keyword': keyword,
        'exact': body.exact,
        'total': len(rows),
        'hits': rows,
    }

class EmergencyCodeSearchIn(BaseModel):
    keyword: str
    case_sensitive: bool = False
    source: str = ''
    repo: str = ''
    limit: int = 200

@app.post('/api/emergency/code-search')
def emergency_code_search(body: EmergencyCodeSearchIn, _: str = Depends(require_auth)):
    keyword = body.keyword.strip()
    if not keyword:
        raise HTTPException(400, '检索关键词不能为空')
    hits = search_synced_code(
        keyword=keyword,
        case_sensitive=body.case_sensitive,
        source=body.source.strip(),
        repo=normalize_repo_name(body.repo),
        limit=body.limit,
    )
    return {'keyword': keyword, 'total': len(hits), 'hits': hits}

# ── 仓库负责人 API ────────────────────────────────────────────────
class RepoOwnerIn(BaseModel):
    repo               : str
    source             : str = ''
    responsible_person : str = ''
    contact            : str = ''

@app.get('/api/repo-owners')
def list_repo_owners(_: str = Depends(require_auth)):
    return db.get_repo_owners()

@app.post('/api/repo-owners')
def save_repo_owner(body: RepoOwnerIn, _: str = Depends(require_auth)):
    if not body.repo.strip():
        raise HTTPException(400, '仓库路径不能为空')
    db.upsert_repo_owner(body.repo.strip(), body.source.strip(), body.responsible_person.strip(), body.contact.strip())
    return {'ok': True}

@app.delete('/api/repo-owners/{oid}')
def remove_repo_owner(oid: int, _: str = Depends(require_auth)):
    db.delete_repo_owner(oid)
    return {'ok': True}

@app.get('/api/repo-owner-map')
def get_repo_owner_map(_: str = Depends(require_auth)):
    raw = db.get_repo_owner_map()
    # Convert tuple keys to string keys for JSON
    return {f"{k[0]}|{k[1]}": v for k, v in raw.items()}

# ── 提示词管理 API ─────────────────────────────────────────────────
class PromptIn(BaseModel):
    name      : str
    category  : str = 'custom'
    extensions: str = ''
    content   : str
    enabled   : bool = True

@app.get('/api/prompts')
def list_prompts(_: str = Depends(require_auth)):
    return db.get_prompts()

@app.post('/api/prompts')
def create_prompt(body: PromptIn, _: str = Depends(require_auth)):
    if not body.name.strip():
        raise HTTPException(400, '名称不能为空')
    if not body.content.strip():
        raise HTTPException(400, '提示词内容不能为空')
    if not db.add_prompt(body.name.strip(), body.category, body.extensions, body.content):
        raise HTTPException(409, '名称已存在')
    return {'ok': True}

@app.put('/api/prompts/{pid}')
def update_prompt(pid: int, body: PromptIn, _: str = Depends(require_auth)):
    if not body.name.strip():
        raise HTTPException(400, '名称不能为空')
    result = db.update_prompt(pid, body.name.strip(), body.category,
                               body.extensions, body.content, body.enabled)
    if result is None:
        raise HTTPException(409, '名称已存在')
    if not result:
        raise HTTPException(404, '提示词不存在')
    return {'ok': True}

@app.delete('/api/prompts/{pid}')
def delete_prompt(pid: int, _: str = Depends(require_auth)):
    if not db.delete_prompt(pid):
        raise HTTPException(400, '内置提示词不可删除，或不存在')
    return {'ok': True}

@app.post('/api/prompts/{pid}/reset')
def reset_prompt(pid: int, _: str = Depends(require_auth)):
    if not db.reset_prompt(pid):
        raise HTTPException(400, '仅内置提示词支持重置')
    return {'ok': True}

# ── 白名单 API ───────────────────────────────────────────────────
class WhitelistIn(BaseModel):
    repo    : str = ''
    filename: str = ''
    title   : str = ''
    reason  : str = ''

@app.get('/api/whitelist')
def list_whitelist(_: str = Depends(require_auth)):
    return db.get_whitelist()

@app.post('/api/whitelist')
def add_whitelist(body: WhitelistIn, operator: str = Depends(require_auth)):
    if not body.repo and not body.filename and not body.title:
        raise HTTPException(400, '至少填写一个匹配条件（仓库、文件名或漏洞标题）')
    wid = db.add_whitelist(body.repo, body.filename, body.title, body.reason, operator)
    syslog.send('info', 'WHITELIST', f'添加白名单 #{wid}: repo={body.repo} file={body.filename} title={body.title}')
    return {'ok': True, 'id': wid}

@app.delete('/api/whitelist/{wid}')
def remove_whitelist(wid: int, operator: str = Depends(require_auth)):
    db.delete_whitelist(wid)
    syslog.send('info', 'WHITELIST', f'删除白名单 #{wid} by {operator}')
    return {'ok': True}

# ── Syslog 配置 API ───────────────────────────────────────────────
@app.get('/api/syslog-config')
def get_syslog(_: str = Depends(require_auth)):
    return db.get_syslog_config()

@app.post('/api/syslog-config')
def save_syslog(body: dict, _: str = Depends(require_auth)):
    db.set_syslog_config(body)
    syslog.reload(db.get_syslog_config())
    return {'ok': True}

@app.post('/api/syslog-test')
def test_syslog(_: str = Depends(require_auth)):
    cfg = db.get_syslog_config()
    if not cfg.get('host', '').strip():
        raise HTTPException(400, '请先填写 Syslog 服务器地址')
    err = syslog.test_send({**cfg, 'enabled': '1'})
    if err:
        raise HTTPException(400, f'发送失败: {err}')
    return {'ok': True, 'msg': f'测试消息已发送至 {cfg["host"]}:{cfg.get("port", 514)}'}

# ── 即时代码分析 API ──────────────────────────────────────────────
class InstantAnalysisIn(BaseModel):
    filename: str
    code    : str
    language: str = ''
    llm_profile_id: Optional[int] = None


class InstantRepoUrlIn(BaseModel):
    repo_url: str
    llm_profile_id: Optional[int] = None


class InstantExportFileIn(BaseModel):
    filename: str
    summary: str = ''
    error: bool = False
    findings: List[dict] = []


class InstantExportIn(BaseModel):
    mode: str = 'upload'
    files: List[InstantExportFileIn] = []


def _resolve_llm_config(llm_profile_id: Optional[int]):
    if llm_profile_id:
        profile = db.get_llm_profile(llm_profile_id)
        if profile:
            return {
                'provider': profile['provider'],
                'model': profile['model'],
                'api_key': profile['api_key'],
                'base_url': profile['base_url'],
            }
    return db.get_llm_config()


_INSTANT_MAX_FILE_SIZE = 300_000    # 单文件最大 300 KB
_INSTANT_MAX_TOTAL_SIZE = 1_000_000 # 总计最大 1 MB
_INSTANT_MAX_FILES = 50             # 最多分析 50 个文件
_INSTANT_SKIP_DIRS = {'.git', '.svn', '.hg', '.idea', '.vscode', 'node_modules', 'dist', 'build', 'target'}


def _instant_prepare_runtime(llm_profile_id: Optional[int]):
    from analyzer import build_llm_caller, _find_prompt, _parse, AUDIT_EXTENSIONS, SKIP_EXTENSIONS
    llm_cfg = _resolve_llm_config(llm_profile_id)
    call_fn = build_llm_caller(llm_cfg)
    if not call_fn:
        raise HTTPException(400, '未配置 LLM，请先在 AI 配置中填写 API Key')
    return {
        'call_fn': call_fn,
        'prompts': db.get_prompts_for_analysis(),
        'find_prompt': _find_prompt,
        'parse': _parse,
        'audit_extensions': AUDIT_EXTENSIONS,
        'skip_extensions': SKIP_EXTENSIONS,
    }


def _instant_should_analyze(fname: str, audit_extensions: set, skip_extensions: set) -> bool:
    _, ext = os.path.splitext(fname)
    ext = ext.lower()
    base = os.path.basename(fname)
    if ext in skip_extensions:
        return False
    return ext in audit_extensions or base in audit_extensions


def _instant_pick_prompt(fname: str, prompts: list, find_prompt):
    prompt_tpl = find_prompt(fname, prompts)
    if not prompt_tpl:
        backend_prompts = [p for p in prompts if p.get('category') == 'backend']
        prompt_tpl = (backend_prompts[0]['content'] if backend_prompts
                      else prompts[0]['content'] if prompts else None)
    return prompt_tpl


def _instant_analyze_collected(collected: list[tuple[str, bytes]], runtime: dict) -> list:
    call_fn = runtime['call_fn']
    prompts = runtime['prompts']
    find_prompt = runtime['find_prompt']
    parse = runtime['parse']

    def _analyze_one(fname: str, data: bytes):
        try:
            code = data.decode('utf-8', errors='replace')
        except Exception:
            return {'filename': fname, 'findings': [], 'summary': '文件解码失败', 'error': True}
        if not code.strip():
            return {'filename': fname, 'findings': [], 'summary': '文件为空'}

        prompt_tpl = _instant_pick_prompt(fname, prompts, find_prompt)
        if not prompt_tpl:
            return {'filename': fname, 'findings': [], 'summary': '未找到可用提示词', 'error': True}

        diff = '\n'.join(f'+{line}' for line in code.splitlines())
        prompt = prompt_tpl.format(filename=fname, message='即时分析', diff=diff)
        try:
            raw_resp = call_fn(prompt)
            result = parse(raw_resp)
            return {
                'filename': fname,
                'findings': result.get('findings', []),
                'summary': result.get('summary', ''),
            }
        except Exception as e:
            return {'filename': fname, 'findings': [], 'summary': f'LLM 分析失败: {e}', 'error': True}

    import concurrent.futures
    results = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as pool:
        futs = {pool.submit(_analyze_one, fn, data): fn for fn, data in collected}
        for fut in concurrent.futures.as_completed(futs):
            results.append(fut.result())
    results.sort(key=lambda r: r['filename'])
    return results


def _extract_repo_host(repo_url: str) -> str:
    repo_url = (repo_url or '').strip()
    if not repo_url:
        return ''
    parsed = urlparse(repo_url)
    if parsed.scheme and parsed.netloc:
        return parsed.netloc.split('@')[-1].split(':')[0].lower()
    m = re.match(r'^[^@]+@([^:]+):', repo_url)
    if m:
        return m.group(1).lower()
    return ''


def _normalize_host(value: str) -> str:
    return (value or '').strip().lower().split('@')[-1].split(':')[0]


def _match_account_for_repo(repo_url: str) -> Optional[dict]:
    host = _extract_repo_host(repo_url)
    rows = db.get_accounts()
    if not rows:
        return None
    for acc in rows:
        acc_host = _extract_repo_host(acc.get('url', ''))
        if host and acc_host and host == _normalize_host(acc_host):
            return acc
    if host == 'github.com':
        return next((a for a in rows if (a.get('platform') or '').lower() == 'github'), None)
    if host == 'gitlab.com':
        return next((a for a in rows if (a.get('platform') or '').lower() == 'gitlab'), None)
    return None


def _inject_token_to_repo_url(repo_url: str, token: str, platform: str) -> str:
    parsed = urlparse(repo_url)
    if not token or parsed.scheme not in ('http', 'https') or not parsed.netloc:
        return repo_url
    host = parsed.netloc.split('@')[-1]
    safe_token = quote(token, safe='')
    plat = (platform or '').lower()
    if plat in ('gitlab', 'tgit', 'codeup'):
        auth = f'oauth2:{safe_token}'
    else:
        auth = safe_token
    return urlunparse((parsed.scheme, f'{auth}@{host}', parsed.path, parsed.params, parsed.query, parsed.fragment))


def _sanitize_clone_error(msg: str, token: str) -> str:
    txt = (msg or '').strip()
    if token:
        txt = txt.replace(token, '***')
        txt = txt.replace(quote(token, safe=''), '***')
    return txt


def _clone_repo_temp(repo_url: str, account: Optional[dict]) -> tuple[str, str]:
    token = (account or {}).get('token', '')
    platform = (account or {}).get('platform', '')
    clone_url = _inject_token_to_repo_url(repo_url, token, platform)
    tmp_dir = tempfile.mkdtemp(prefix='instant_repo_')
    repo_dir = os.path.join(tmp_dir, 'repo')
    proc = subprocess.run(
        ['git', 'clone', '--depth', '1', clone_url, repo_dir],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        timeout=180,
    )
    if proc.returncode == 0:
        return tmp_dir, repo_dir

    stderr = _sanitize_clone_error(proc.stderr, token).lower()
    if any(k in stderr for k in ('could not resolve host', 'failed to connect', 'connection timed out', 'operation timed out', 'network is unreachable', 'name or service not known', 'connection refused')):
        shutil.rmtree(tmp_dir, ignore_errors=True)
        raise HTTPException(502, '仓库拉取失败：网络不通或仓库不可达，请检查网络连通性和仓库地址')
    if any(k in stderr for k in ('authentication failed', 'could not read username', 'access denied', '403', 'permission denied', 'repository not found')):
        shutil.rmtree(tmp_dir, ignore_errors=True)
        raise HTTPException(401, '仓库拉取失败：认证失败或无权限，请检查代码库配置中的 Token')

    shutil.rmtree(tmp_dir, ignore_errors=True)
    msg = _sanitize_clone_error(proc.stderr or proc.stdout, token).strip()
    brief = msg.splitlines()[-1] if msg else '未知错误'
    raise HTTPException(400, f'仓库拉取失败：{brief}')


def _build_instant_markdown(files: list) -> str:
    sev_zh = {'critical': '严重', 'high': '高危', 'medium': '中危', 'low': '低危'}
    lines = [
        '# 即时分析报告',
        '',
        f'> 生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}',
        ''
    ]
    total = 0
    for f in files:
        findings = f.get('findings', []) or []
        total += len(findings)
        lines.append(f'## 文件：{f.get("filename", "")}')
        if f.get('error'):
            lines.append(f'- 状态：失败')
            lines.append(f'- 原因：{f.get("summary", "")}')
            lines.append('')
            continue
        lines.append(f'- 发现问题：{len(findings)}')
        if f.get('summary'):
            lines.append(f'- 摘要：{f.get("summary", "")}')
        lines.append('')
        for i, item in enumerate(findings, 1):
            sev = sev_zh.get(item.get('severity', 'low'), item.get('severity', 'low'))
            lines += [
                f'### {i}. [{sev}] {item.get("title", "")}',
                f'- 位置：`{item.get("filename") or f.get("filename","")}`'
                + (f' 第 {item.get("line")} 行' if item.get('line') else ''),
                f'- 描述：{item.get("description", "")}',
            ]
            if item.get('recommendation'):
                lines.append(f'- 修复建议：{item.get("recommendation", "")}')
            lines.append('')
    lines.insert(3, f'> 文件数：{len(files)}，问题总数：{total}')
    return '\n'.join(lines)


def _build_instant_html(files: list) -> str:
    sev_color = {'critical': '#ff4d4f', 'high': '#fa8c16', 'medium': '#d4b106', 'low': '#52c41a'}
    sev_zh = {'critical': '严重', 'high': '高危', 'medium': '中危', 'low': '低危'}
    total = sum(len(f.get('findings', []) or []) for f in files)
    parts = ["""<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1"><title>即时分析报告</title>
<style>body{font-family:-apple-system,BlinkMacSystemFont,'PingFang SC',sans-serif;background:#f5f5f5;color:#1a1a1a;padding:20px}
.c{max-width:960px;margin:0 auto}.card{background:#fff;border-radius:10px;padding:14px 16px;margin-bottom:12px;box-shadow:0 1px 4px rgba(0,0,0,.06)}
.sev{font-size:11px;color:#fff;padding:2px 8px;border-radius:4px;font-weight:700}.title{font-size:14px;font-weight:700}.meta{font-size:12px;color:#888}
.desc{font-size:13px;color:#555;line-height:1.6}.rec{font-size:12px;color:#1677ff;background:#f0f5ff;border-radius:6px;padding:8px 10px}
</style></head><body><div class="c">"""]
    parts.append(f"<h2>即时分析报告</h2><div class='meta'>生成时间：{_html.escape(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}；文件数：{len(files)}；问题总数：{total}</div><br>")
    for f in files:
        fn = _html.escape(str(f.get('filename', '')))
        summary = _html.escape(str(f.get('summary', '')))
        findings = f.get('findings', []) or []
        parts.append(f"<div class='card'><div class='title'>📄 {fn}</div><div class='meta'>发现问题：{len(findings)}</div>")
        if f.get('error'):
            parts.append(f"<div class='meta' style='color:#ff4d4f'>分析失败：{summary}</div></div>")
            continue
        if summary:
            parts.append(f"<div class='meta'>摘要：{summary}</div>")
        for item in findings:
            sev = item.get('severity', 'low')
            parts.append("<hr style='border:none;border-top:1px solid #f0f0f0;margin:10px 0'>")
            parts.append(
                f"<span class='sev' style='background:{sev_color.get(sev, '#888')}'>{_html.escape(sev_zh.get(sev, sev))}</span> "
                f"<span class='title'>{_html.escape(str(item.get('title','')))}</span>"
            )
            line = f" 第 {_html.escape(str(item.get('line')))} 行" if item.get('line') else ''
            parts.append(f"<div class='meta'>位置：{_html.escape(str(item.get('filename') or f.get('filename','')))}{line}</div>")
            parts.append(f"<div class='desc'>{_html.escape(str(item.get('description','')))}</div>")
            if item.get('recommendation'):
                parts.append(f"<div class='rec'>修复建议：{_html.escape(str(item.get('recommendation')))}</div>")
        parts.append("</div>")
    parts.append("</div></body></html>")
    return ''.join(parts)


def _build_instant_docx(files: list) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading('即时分析报告', 0)
    total = sum(len(f.get('findings', []) or []) for f in files)
    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'文件数：{len(files)}，问题总数：{total}')
    sev_zh = {'critical': '严重', 'high': '高危', 'medium': '中危', 'low': '低危'}
    for f in files:
        doc.add_heading(f'文件：{f.get("filename", "")}', level=1)
        if f.get('error'):
            doc.add_paragraph(f'分析失败：{f.get("summary", "")}')
            continue
        if f.get('summary'):
            doc.add_paragraph(f'摘要：{f.get("summary", "")}')
        findings = f.get('findings', []) or []
        if not findings:
            doc.add_paragraph('未发现安全问题')
            continue
        for i, item in enumerate(findings, 1):
            sev = sev_zh.get(item.get('severity', 'low'), item.get('severity', 'low'))
            doc.add_heading(f'{i}. [{sev}] {item.get("title", "")}', level=2)
            loc = f'{item.get("filename") or f.get("filename","")}'
            if item.get('line'):
                loc += f' 第 {item.get("line")} 行'
            doc.add_paragraph(f'位置：{loc}')
            doc.add_paragraph(f'描述：{item.get("description", "")}')
            if item.get('recommendation'):
                doc.add_paragraph(f'修复建议：{item.get("recommendation", "")}')
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_instant_pdf(files: list) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4
    y = h - 36
    c.setFont('STSong-Light', 12)

    def write_line(txt: str):
        nonlocal y
        if y < 40:
            c.showPage()
            c.setFont('STSong-Light', 12)
            y = h - 36
        c.drawString(36, y, txt[:120])
        y -= 16

    total = sum(len(f.get('findings', []) or []) for f in files)
    write_line('即时分析报告')
    write_line(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    write_line(f'文件数：{len(files)}，问题总数：{total}')
    write_line('')
    sev_zh = {'critical': '严重', 'high': '高危', 'medium': '中危', 'low': '低危'}
    for f in files:
        write_line(f'文件：{f.get("filename", "")}')
        if f.get('error'):
            write_line(f'分析失败：{f.get("summary", "")}')
            write_line('')
            continue
        if f.get('summary'):
            write_line(f'摘要：{f.get("summary", "")}')
        findings = f.get('findings', []) or []
        if not findings:
            write_line('未发现安全问题')
            write_line('')
            continue
        for i, item in enumerate(findings, 1):
            sev = sev_zh.get(item.get('severity', 'low'), item.get('severity', 'low'))
            write_line(f'{i}. [{sev}] {item.get("title", "")}')
            loc = f'{item.get("filename") or f.get("filename","")}'
            if item.get('line'):
                loc += f' 第 {item.get("line")} 行'
            write_line(f'位置：{loc}')
            if item.get('description'):
                write_line(f'描述：{item.get("description", "")}')
            if item.get('recommendation'):
                write_line(f'建议：{item.get("recommendation", "")}')
            write_line('')
    c.save()
    return buf.getvalue()

@app.post('/api/analyze/instant')
def instant_analyze(body: InstantAnalysisIn, _: str = Depends(require_auth)):
    """对任意代码片段直接进行 LLM 安全分析，无需 git commit。"""
    _require_license_if_enabled('instant_analysis')
    from analyzer import build_llm_caller, _find_prompt, _parse, AUDIT_EXTENSIONS, SKIP_EXTENSIONS
    import os as _os

    filename = body.filename.strip() or 'snippet.py'
    code     = body.code.strip()
    if not code:
        raise HTTPException(400, '代码内容不能为空')
    if len(code) > 30000:
        raise HTTPException(400, '代码内容过长，最多 30000 字符')

    llm_cfg = _resolve_llm_config(body.llm_profile_id)
    call_fn = build_llm_caller(llm_cfg)
    if not call_fn:
        raise HTTPException(400, '未配置 LLM，请先在 AI 配置中填写 API Key')

    prompts = db.get_prompts_for_analysis()
    prompt_tpl = _find_prompt(filename, prompts)
    if not prompt_tpl:
        # fallback to backend prompt
        backend_prompts = [p for p in prompts if p.get('category') == 'backend']
        prompt_tpl = backend_prompts[0]['content'] if backend_prompts else prompts[0]['content'] if prompts else None
    if not prompt_tpl:
        raise HTTPException(500, '未找到可用提示词')

    # 构造一个"全量 diff"形式（+行前缀）
    diff = '\n'.join(f'+{line}' for line in code.splitlines())
    prompt = prompt_tpl.format(filename=filename, message='即时分析', diff=diff)
    try:
        raw    = call_fn(prompt)
        result = _parse(raw)
        return {
            'filename': filename,
            'findings': result.get('findings', []),
            'summary' : result.get('summary', ''),
        }
    except Exception as e:
        raise HTTPException(500, f'LLM 分析失败: {e}')

@app.post('/api/analyze/instant-upload')
async def instant_analyze_upload(
    files: List[UploadFile] = File(...),
    llm_profile_id: Optional[int] = Form(None),
    credentials: HTTPAuthorizationCredentials = Depends(_http_bearer)
):
    """上传文件/文件夹/压缩包进行 LLM 安全分析。"""
    _ = require_auth(credentials)
    _require_license_if_enabled('instant_analysis')
    runtime = _instant_prepare_runtime(llm_profile_id)

    # 收集 (filename, content_bytes) 列表
    collected: list[tuple[str, bytes]] = []
    total_bytes = 0
    _MAX_ZIP_ENTRY_SIZE = _INSTANT_MAX_FILE_SIZE
    _MAX_ZIP_RATIO = 200  # 解压比阈值，避免 zip bomb

    for upload in files:
        raw = await upload.read()
        fname = upload.filename or 'file'

        # zip 压缩包：展开
        _, ext = os.path.splitext(fname)
        if ext.lower() == '.zip':
            try:
                with zipfile.ZipFile(io.BytesIO(raw)) as zf:
                    for entry in zf.infolist():
                        if entry.is_dir():
                            continue
                        entry_name = entry.filename.lstrip('/')
                        if not _instant_should_analyze(entry_name, runtime['audit_extensions'], runtime['skip_extensions']):
                            continue
                        if entry.file_size > _MAX_ZIP_ENTRY_SIZE:
                            continue
                        if entry.compress_size > 0 and (entry.file_size / max(1, entry.compress_size)) > _MAX_ZIP_RATIO:
                            continue
                        if total_bytes + max(0, entry.file_size) > _INSTANT_MAX_TOTAL_SIZE:
                            break
                        data = zf.read(entry)
                        if len(data) > _INSTANT_MAX_FILE_SIZE:
                            continue
                        collected.append((entry_name, data))
                        total_bytes += len(data)
                        if total_bytes > _INSTANT_MAX_TOTAL_SIZE or len(collected) >= _INSTANT_MAX_FILES:
                            break
            except zipfile.BadZipFile:
                raise HTTPException(400, f'{fname} 不是有效的 ZIP 文件')
        else:
            if not _instant_should_analyze(fname, runtime['audit_extensions'], runtime['skip_extensions']):
                continue
            if len(raw) > _INSTANT_MAX_FILE_SIZE:
                continue
            collected.append((fname, raw))
            total_bytes += len(raw)

        if total_bytes > _INSTANT_MAX_TOTAL_SIZE or len(collected) >= _INSTANT_MAX_FILES:
            break

    if not collected:
        raise HTTPException(400, '未找到可分析的代码文件（不支持的类型或文件过大）')

    results = _instant_analyze_collected(collected, runtime)
    total_findings = sum(len(r['findings']) for r in results)
    return {'files': results, 'total_findings': total_findings}


@app.post('/api/analyze/instant-repo-url')
def instant_analyze_repo_url(body: InstantRepoUrlIn, _: str = Depends(require_auth)):
    """输入仓库 URL，拉取后进行即时安全分析。"""
    _require_license_if_enabled('instant_analysis')
    repo_url = (body.repo_url or '').strip()
    if not repo_url:
        raise HTTPException(400, '仓库 URL 不能为空')
    if not (repo_url.startswith('http://') or repo_url.startswith('https://') or re.match(r'^[^@]+@[^:]+:.+', repo_url)):
        raise HTTPException(400, '仓库 URL 格式不正确，仅支持 HTTP(S) 或 SSH 地址')

    runtime = _instant_prepare_runtime(body.llm_profile_id)
    account = _match_account_for_repo(repo_url)

    tmp_dir = ''
    try:
        tmp_dir, repo_dir = _clone_repo_temp(repo_url, account)
        collected: list[tuple[str, bytes]] = []
        total_bytes = 0
        for root, dirs, files in os.walk(repo_dir):
            dirs[:] = [d for d in dirs if d not in _INSTANT_SKIP_DIRS]
            for name in files:
                full = os.path.join(root, name)
                if os.path.islink(full):
                    continue
                rel = os.path.relpath(full, repo_dir).replace('\\', '/')
                if not _instant_should_analyze(rel, runtime['audit_extensions'], runtime['skip_extensions']):
                    continue
                try:
                    size = os.path.getsize(full)
                    if size > _INSTANT_MAX_FILE_SIZE:
                        continue
                    with open(full, 'rb') as fh:
                        data = fh.read()
                except Exception:
                    continue
                collected.append((rel, data))
                total_bytes += len(data)
                if total_bytes > _INSTANT_MAX_TOTAL_SIZE or len(collected) >= _INSTANT_MAX_FILES:
                    break
            if total_bytes > _INSTANT_MAX_TOTAL_SIZE or len(collected) >= _INSTANT_MAX_FILES:
                break
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    if not collected:
        raise HTTPException(400, '未找到可分析的代码文件（不支持的类型、文件过大或仓库为空）')

    results = _instant_analyze_collected(collected, runtime)
    total_findings = sum(len(r['findings']) for r in results)
    return {'files': results, 'total_findings': total_findings}


@app.post('/api/analyze/instant/export')
def export_instant_report(body: InstantExportIn, format: str = 'html', _: str = Depends(require_auth)):
    if format not in ('html', 'pdf', 'docx', 'markdown'):
        raise HTTPException(400, '无效格式，可选：html / pdf / docx / markdown')
    files = [f.model_dump() for f in (body.files or [])]
    if not files:
        raise HTTPException(400, '暂无可导出的分析结果')
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')

    if format == 'html':
        content = _build_instant_html(files)
        return Response(
            content=content.encode('utf-8'),
            media_type='text/html; charset=utf-8',
            headers={'Content-Disposition': f'attachment; filename="instant_report_{ts}.html"'}
        )
    if format == 'markdown':
        content = _build_instant_markdown(files)
        return Response(
            content=content.encode('utf-8'),
            media_type='text/markdown; charset=utf-8',
            headers={'Content-Disposition': f'attachment; filename="instant_report_{ts}.md"'}
        )
    if format == 'docx':
        content = _build_instant_docx(files)
        return Response(
            content=content,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="instant_report_{ts}.docx"'}
        )
    content = _build_instant_pdf(files)
    return Response(
        content=content,
        media_type='application/pdf',
        headers={'Content-Disposition': f'attachment; filename="instant_report_{ts}.pdf"'}
    )


@app.get('/api/status')
def status(_: str = Depends(require_auth)):
    scanning = not _scan_lock.acquire(blocking=False)
    if not scanning:
        _scan_lock.release()

    schedules_out = []
    for s in db.get_scan_schedules():
        job = scheduler.get_job(f"sched_{s['id']}")
        schedules_out.append({
            'id'      : s['id'],
            'type'    : s['type'],
            'hour'    : s['hour'],
            'minute'  : s['minute'],
            'weekday' : s.get('weekday'),
            'enabled' : bool(s['enabled']),
            'label'   : s.get('label', ''),
            'next_run': str(job.next_run_time) if job else None,
        })

    return {
        'scheduler_running': scheduler.running,
        'scanning'         : scanning,
        'paused'           : scanning and not _pause_event.is_set(),
        'progress'         : _get_scan_progress() if scanning else None,
        'schedules'        : schedules_out,
        'license'          : _get_license_status(),
    }

if __name__ == '__main__':
    import uvicorn
    uvicorn.run('app:app', host='0.0.0.0', port=8000, reload=False)
